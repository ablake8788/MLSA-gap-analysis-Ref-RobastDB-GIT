# docx_renderer.py          # markdown_to_docx

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from tga_cli.domain.models import RunContext, RunArtifacts
from tga_cli.renderers.references import append_references_appendix


def _set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.08


def _add_cover(doc: Document, title: str, ctx: RunContext) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)

    meta = [
        ("Competitor", ctx.competitor_url),
        ("Generated", ctx.generated_at),
        ("Input file", ctx.input_file.name),
    ]
    if ctx.baseline_enabled:
        meta.append(("Secondary Competitor", ctx.baseline_url))

    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = str(k)
        table.cell(i, 1).text = str(v)

    doc.add_paragraph("")


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    if not cells:
        return False
    for c in cells:
        c2 = c.replace(":", "").replace("-", "")
        if c2 != "":
            return False
    return True


def _split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _render_md_table(doc: Document, table_lines: list[str]) -> None:
    header = _split_table_row(table_lines[0])
    rows = [_split_table_row(x) for x in table_lines[2:] if _is_table_row(x)]
    cols = len(header)

    t = doc.add_table(rows=1 + len(rows), cols=cols)
    t.style = "Light Shading Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for j, h in enumerate(header):
        cell = t.cell(0, j)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, rvals in enumerate(rows, start=1):
        rvals = (rvals + [""] * cols)[:cols]
        for j, v in enumerate(rvals):
            t.cell(i, j).text = v


class DocxRenderer:
    def render(self, *, report_md: str, ctx: RunContext, artifacts: RunArtifacts) -> None:
        doc = Document()
        _set_doc_defaults(doc)
        _add_cover(doc, "Competitors Gap Analysis Report", ctx)

        lines = report_md.splitlines()
        i = 0
        in_code = False
        code_buf: list[str] = []

        def flush_code() -> None:
            nonlocal code_buf
            if not code_buf:
                return
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            code_buf = []

        while i < len(lines):
            raw = lines[i].rstrip("\n")
            s = raw.strip()

            if s.startswith("```"):
                in_code = not in_code
                if not in_code:
                    flush_code()
                i += 1
                continue

            if in_code:
                code_buf.append(raw)
                i += 1
                continue

            if _is_table_row(raw) and (i + 1) < len(lines) and _is_table_sep(lines[i + 1]):
                table_block = [raw, lines[i + 1]]
                i += 2
                while i < len(lines) and _is_table_row(lines[i]):
                    table_block.append(lines[i])
                    i += 1
                _render_md_table(doc, table_block)
                continue

            if not s:
                doc.add_paragraph("")
                i += 1
                continue

            if s.startswith("# "):
                doc.add_heading(s[2:].strip(), level=1)
            elif s.startswith("## "):
                doc.add_heading(s[3:].strip(), level=2)
            elif s.startswith("### "):
                doc.add_heading(s[4:].strip(), level=3)
            elif s.startswith(("- ", "* ")):
                doc.add_paragraph(s[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(raw)

            i += 1

        flush_code()

        import logging
        logger = logging.getLogger("tga_cli")

        # References appendix (enabled by preset/config; safe defaults if fields missing)
        include_refs = bool(getattr(ctx, "include_references", True))
        ref_format = (getattr(ctx, "reference_format", None) or "").lower()
        logger.info("DOCX references enabled=%s format=%s", include_refs, ref_format)


        if include_refs and ref_format == "appendix":
            append_references_appendix(
                doc,
                competitor_url=ctx.competitor_url,
                baseline_url=ctx.baseline_url if ctx.baseline_enabled else None,
                input_file=str(ctx.input_file),
                generated_at=ctx.generated_at,
                competitor_text=getattr(ctx, "competitor_text", None),
                baseline_text=getattr(ctx, "baseline_text", None),
                model_output_text=report_md,
            )

        doc.save(str(artifacts.docx_path))



